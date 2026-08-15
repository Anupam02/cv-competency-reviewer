from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from pypdf import PdfReader
from docx import Document


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class IngestedDocument:
    text: str
    filename: str
    media_type: str


def ingest_path(path: str | Path) -> IngestedDocument:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type {suffix}. Use PDF, DOCX, TXT, or MD.")
    data = path.read_bytes()
    return ingest_bytes(data, filename=path.name, suffix=suffix)


def ingest_bytes(data: bytes, filename: str, suffix: str | None = None) -> IngestedDocument:
    suffix = (suffix or Path(filename).suffix).lower()
    if suffix == ".pdf":
        text = _pdf_text(data)
        media = "application/pdf"
    elif suffix == ".docx":
        text = _docx_text(data)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace")
        media = "text/plain"
    else:
        raise ValueError(f"Unsupported file type {suffix}. Use PDF, DOCX, TXT, or MD.")
    cleaned = _normalise(text)
    if not cleaned.strip():
        raise ValueError("No extractable text was found in the CV.")
    return IngestedDocument(text=cleaned, filename=filename, media_type=media)


def _pdf_text(data: bytes) -> str:
    from io import BytesIO

    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _docx_text(data: bytes) -> str:
    from io import BytesIO

    document = Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _normalise(text: str) -> str:
    text = text.replace("\x00", " ")
    lines = [line.rstrip() for line in text.splitlines()]
    collapsed: list[str] = []
    blank = 0
    for line in lines:
        if not line.strip():
            blank += 1
            if blank <= 1:
                collapsed.append("")
            continue
        blank = 0
        collapsed.append(line)
    return "\n".join(collapsed).strip()
